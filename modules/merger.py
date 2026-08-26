
import os
import io
import pandas as pd
import docx
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

class DocumentMerger:
    """
    Class responsible for merging multiple processed documents using 2 distinct strategies
    (or skipping) and exporting the output to standard local files (.docx, .xlsx, and .pdf).
    """

    def __init__(self):
        pass

    def merge_documents(self, documents_dict: dict, merge_mode: str) -> str:
        """
        Merges text from multiple documents based on the chosen strategy:
        - Option 1: Sequential Concatenation.
        - Option 2: Grouped by Document Type.
        - Option 3: Skip Merge (Raw text return).
        """
        if merge_mode == "Nenhum (Apenas Análise Direta)":
            # Retorna o conteúdo bruto sem formatação extra de merge
            return "\n\n".join([f"[Source: {filename}]\n{text}" for filename, text in documents_dict.items()])

        combined_text = ""

        if merge_mode == "Option 1: Sequential Concatenation":
            for filename, text in documents_dict.items():
                combined_text += f"\n\n========================================\n"
                combined_text += f"FILE SOURCE: {filename}\n"
                combined_text += f"========================================\n\n"
                combined_text += text

        elif merge_mode == "Option 2: Grouped by Document Type":
            text_docs = {}
            data_docs = {}
            for filename, text in documents_dict.items():
                if filename.endswith(('.xlsx', '.csv', '.xls')):
                    data_docs[filename] = text
                else:
                    text_docs[filename] = text

            combined_text += "--- SECTION A: TEXTUAL DOCUMENTS ---\n"
            for filename, text in text_docs.items():
                combined_text += f"\n[Source: {filename}]\n{text}\n"

            combined_text += "\n\n--- SECTION B: TABULAR/SPREADSHEET DOCUMENTS ---\n"
            for filename, text in data_docs.items():
                combined_text += f"\n[Source: {filename}]\n{text}\n"

        return combined_text

    def export_to_docx(self, content: str) -> io.BytesIO:
        """Exports the summarized or merged text into a basic formatted Word (.docx) file."""
        doc = docx.Document()
        doc.add_heading("FrogPDF - Document Report", 0)
        
        for paragraph in content.split("\n"):
            if paragraph.startswith("# "):
                doc.add_heading(paragraph.replace("# ", ""), level=1)
            elif paragraph.startswith("## "):
                doc.add_heading(paragraph.replace("## ", ""), level=2)
            elif paragraph.strip():
                doc.add_paragraph(paragraph)
                
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def export_to_pdf(self, content: str) -> io.BytesIO:
        """Exports the analyzed text content into a clean formatted PDF report."""
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        story = []
        
        styles = getSampleStyleSheet()
        normal_style = styles['Normal']
        title_style = styles['Heading1']
        
        story.append(Paragraph("FrogPDF - Relatório de Análise Documental", title_style))
        story.append(Spacer(1, 12))
        
        for line in content.split("\n"):
            if line.strip():
                story.append(Paragraph(line, normal_style))
                story.append(Spacer(1, 4))
                
        doc.build(story)
        buffer.seek(0)
        return buffer

    def export_markdown_tables_to_excel(self, markdown_text: str) -> io.BytesIO:
        """Basic parser to find markdown tables in the LLM response and export them into an Excel (.xlsx) file."""
        buffer = io.BytesIO()
        try:
            lines = markdown_text.split("\n")
            table_lines = [line for line in lines if "|" in line]
            
            if len(table_lines) > 2:
                cleaned_lines = [l for l in table_lines if "---" not in l]
                data = []
                for line in cleaned_lines:
                    cols = [c.strip() for c in line.split("|")[1:-1]]
                    data.append(cols)
                
                if len(data) > 1:
                    df = pd.DataFrame(data[1:], columns=data[0])
                    with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                        df.to_excel(writer, index=False, sheet_name="Extracted_Table")
                    buffer.seek(0)
                    return buffer
        except Exception:
            pass
            
        df = pd.DataFrame({"Info": ["No structured tables found to export."]})
        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name="Empty")
        buffer.seek(0)
        return buffer